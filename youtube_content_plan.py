
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Page Setup ---
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# --- Styles helpers ---
def set_heading(para, text, level=1, color_hex="1A1A2E"):
    para.clear()
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    elif level == 3:
        run.font.size = Pt(13)
    r, g, b = bytes.fromhex(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    para.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    para.paragraph_format.space_after = Pt(6)

def add_heading(doc, text, level=1, color_hex="1A1A2E"):
    para = doc.add_paragraph()
    set_heading(para, text, level, color_hex)
    return para

def add_body(doc, text, bold=False, color_hex=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    if color_hex:
        r, g, b = bytes.fromhex(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    para.paragraph_format.space_after = Pt(4)
    return para

def add_bullet(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = para.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run2 = para.add_run(text)
        run2.font.size = Pt(11)
    else:
        run = para.add_run(text)
        run.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(3)
    return para

def add_divider(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(8)
    return para

def add_video_idea(doc, num, emoji, title, why, format_, length, hook):
    p = doc.add_paragraph()
    run = p.add_run(f"{emoji} Video Idea #{num}: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x10, 0x69, 0xAD)
    run2 = p.add_run(title)
    run2.bold = True
    run2.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    
    add_bullet(doc, why, "Why it works: ")
    add_bullet(doc, format_, "Format: ")
    add_bullet(doc, length, "Length: ")
    add_bullet(doc, hook, "Hook: ")

# ============================================================
# COVER PAGE
# ============================================================
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("🎬 YouTube Growth Master Plan")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x10, 0x69, 0xAD)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("4-Channel Content Strategy | March 2026")
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

owner = doc.add_paragraph()
owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
owner.paragraph_format.space_before = Pt(10)
r3 = owner.add_run("Dashon McFarlane")
r3.font.size = Pt(13)
r3.bold = True

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = date_p.add_run(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}")
r4.font.size = Pt(11)
r4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

channels_box = doc.add_paragraph()
channels_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
channels_box.paragraph_format.space_before = Pt(24)
r5 = channels_box.add_run(
    "🌿 NaturalHealingRevival   •   💰 Richer By Design\n"
    "🌍 Aria's Great Adventures   •   🌌 Dashon McFarlane"
)
r5.font.size = Pt(12)
r5.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
add_heading(doc, "📋 Executive Summary", 1, "1069AD")
add_body(doc, 
    "This content strategy report covers all 4 of your YouTube channels, each starting from 0 subscribers. "
    "Based on live research conducted in March 2026, this plan provides channel-specific trending topics, "
    "optimized video titles, best formats for growth, and a 90-day launch playbook.")
add_body(doc, "")

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Channel"
hdr[1].text = "Niche"
hdr[2].text = "CPM Range"
hdr[3].text = "Best Format"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)

rows_data = [
    ("🌿 NaturalHealingRevival", "Natural Health & Wellness", "$15–$40+", "Shorts funnel → Long-form"),
    ("💰 Richer By Design", "Personal Finance & Wealth", "$12–$22", "Long-form talking head + Shorts"),
    ("🌍 Aria's Great Adventures", "Kids & Family Adventure", "$5–$12", "Live-action vlog + daily Shorts"),
    ("🌌 Dashon McFarlane", "Personal Brand / Marketing", "$8–$18", "Talking head + Shorts funnel"),
]
for i, (ch, niche, cpm, fmt) in enumerate(rows_data):
    row = table.rows[i+1].cells
    row[0].text = ch
    row[1].text = niche
    row[2].text = cpm
    row[3].text = fmt
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
add_body(doc, "Universal Growth Strategy (All Channels): Post 3–5 Shorts/week as a discovery funnel into 1–2 long-form videos/week. Consistency + SEO-optimized titles beat viral chasing every time.", color_hex="444444")

add_divider(doc)
doc.add_page_break()

# ============================================================
# CHANNEL 1: NATURAL HEALING REVIVAL
# ============================================================
add_heading(doc, "🌿 Channel 1: NaturalHealingRevival", 1, "2E7D32")
add_heading(doc, "Niche: Natural Health, Herbal Remedies & Holistic Wellness", 2, "388E3C")

add_heading(doc, "🔥 Top Trends Right Now (March 2026)", 2, "2E7D32")
trends1 = [
    ("Adaptogens & Stress Herbs", "Ashwagandha, Rhodiola, Holy Basil, Lion's Mane are exploding. Search demand for cortisol/burnout solutions is at an all-time high."),
    ("Gut Microbiome & Digestive Healing", "Gut-brain axis content generates millions of views. \"Healing your gut naturally\" and fermented food protocols dominate."),
    ("Natural Longevity & Anti-Aging", "Biohacking Lite — audiences want NAD+ boosters, red light therapy, and ancestral diets WITHOUT expensive lab tests."),
    ("Liver Detox & Mineral Health", "Viral Shorts series on copper/liver detox spreading rapidly. Castor oil, dandelion root, milk thistle spiking in search."),
    ("'Doctors vs. Nature' Controversy Content", "Barbara O'Neill-style content ('What Doctors Won't Tell You') drives enormous organic discovery on Shorts and search."),
    ("Metabolic Health & Blood Sugar — Naturally", "Natural insulin sensitivity and anti-inflammatory content. Overlaps with anti-GLP-1 drug narrative gaining momentum."),
]
for name, desc in trends1:
    add_bullet(doc, f"{desc}", f"{name}: ")

add_heading(doc, "🎬 Best Format", 2, "2E7D32")
add_bullet(doc, "Post 3–5 Shorts/week with shocking hooks (e.g. 'This 1 herb lowers cortisol in 3 days')")
add_bullet(doc, "Each Short drives to a full long-form video — health CPM is $15–$40+ per 1,000 views")
add_bullet(doc, "Talking head builds expert trust; voiceover + B-roll works for listicle/faceless channels")
add_bullet(doc, "SEO title formula: [Specific Benefit] + [Time Frame] + [Authority Signal or Curiosity Trigger]")

add_heading(doc, "🏆 Top 5 Video Ideas", 2, "2E7D32")
add_video_idea(doc, 1, "🌿", 
    '"The Ancient Herb That Lowers Cortisol in 7 Days (Doctors Are Finally Admitting It)"',
    "Taps the cortisol/stress epidemic; 'doctors finally admitting' triggers curiosity; '7 days' gives a concrete promise.",
    "Talking head + supplement visuals, 10–14 min",
    "10–14 minutes",
    '"Most people are walking around with cortisol levels 3x higher than they should be. This herb has been used for 3,000 years — and science is finally catching up."'
)
add_video_idea(doc, 2, "🦠",
    '"I Healed My Gut in 30 Days Using Only These 4 Forgotten Remedies (No Supplements)"',
    "'30 days' transformation hook; 'forgotten remedies' = natural wisdom angle; 'no supplements' attracts budget-conscious viewers.",
    "Story-driven talking head or voiceover + B-roll, 12–16 min",
    "12–16 minutes",
    '"I spent thousands on probiotics and gut tests. Then I found these 4 remedies that cost almost nothing."'
)
add_video_idea(doc, 3, "⏳",
    '"Scientists Discovered This Herb Reverses Aging at the Cellular Level (And It Costs $3)"',
    "Science credibility + natural angle + affordability. Longevity content growing 40%+ YoY.",
    "Voiceover with graphics + talking head outro, 10–15 min",
    "10–15 minutes",
    '"Longevity researchers at a top university just published findings on this — and you can find it at any grocery store for under $3."'
)
add_video_idea(doc, 4, "🫀",
    '"Do This Every Morning for 7 Days and Watch What Happens to Your Liver (Ancient Remedy)"',
    "Challenge-style format drives retention and return viewers. Liver detox is viral in 2026 Shorts. 'Ancient remedy' = wisdom-based.",
    "Challenge/series format, talking head, 8–12 min",
    "8–12 minutes",
    '"Your liver processes 500+ functions every day. Most people are unknowingly destroying it. Here\'s what to do instead."'
)
add_video_idea(doc, 5, "⚡",
    '"STOP Taking These 5 \'Healthy\' Supplements — They\'re Making Your Inflammation WORSE"',
    "Pattern-interrupt title challenges assumed behavior. Controversy angle drives shares and comments. Massive search intent around inflammation.",
    "Talking head, direct/authoritative tone, 12–15 min",
    "12–15 minutes",
    '"I know this is going to be controversial. But after studying the research, I can\'t stay silent anymore."'
)

add_heading(doc, "📈 90-Day Growth Tips", 2, "2E7D32")
add_bullet(doc, "Claim a micro-identity: 'The Herb Channel,' 'Natural Liver Health,' or 'Ancestral Healing' outperforms broad 'wellness'")
add_bullet(doc, "Post Shorts daily for the first 90 days — use trending audio + 'did you know this herb…' hooks")
add_bullet(doc, "Launch a '30-Day Herbal Reset Challenge' series to retain viewers across multiple videos")
add_bullet(doc, "High-CPM keywords: liver detox, cortisol, longevity herbs, gut microbiome, adaptogens for stress")
add_bullet(doc, "Personal health data stories ('I tested my cortisol and fixed it naturally') are surging in 2026")

add_divider(doc)
doc.add_page_break()

# ============================================================
# CHANNEL 2: RICHER BY DESIGN
# ============================================================
add_heading(doc, "💰 Channel 2: Richer By Design", 1, "1565C0")
add_heading(doc, "Niche: Personal Finance, Wealth Building & Intentional Wealth Design", 2, "1976D2")

add_heading(doc, "🔥 Top Trends Right Now (March 2026)", 2, "1565C0")
trends2 = [
    ("Tariff Economy Anxiety + Wealth Protection", "Trump's 2026 tariffs averaging $700 tax increase per household. Massive search demand for recession-proofing, inflation hedges, and tariff impact on personal finances."),
    ("AI + Wealth Overlap", "Finance content combining AI investment strategy is getting outsized reach. 'Is the AI bubble about to burst? How do I invest in AI?' dominating search."),
    ("Interest Rates & The 'Money Reset' Mindset", "Fed rate cuts expected March 2026 — content around locking in high-yield savings, CD ladders, and bond strategies is spiking."),
    ("Anti-Passive-Income Backlash", "Viral content debunking 'easy passive income' gurus. Authentic, results-oriented creators are winning audience trust."),
    ("'Wealth Design' & Financial Identity", "Micro-niche around designing your financial life intentionally — goal-based investing, lifestyle architecture — is underserved and growing."),
    ("Generational Finance (Gen Z & Millennials)", "Content for Gen Z navigating housing unaffordability, student debt, side hustles vs. careers is surging organically."),
]
for name, desc in trends2:
    add_bullet(doc, f"{desc}", f"{name}: ")

add_heading(doc, "🎬 Best Format", 2, "1565C0")
add_bullet(doc, "Lead with long-form talking head (10–18 min) — finance viewers trust faces; personal brand is part of 'designed wealth' identity")
add_bullet(doc, "Use Shorts as teasers — repurpose 45–60 second clips as hooks posted 24hrs before the long-form")
add_bullet(doc, "Finance is YouTube's highest-paying niche — target $12–$22 CPM with long-form; Shorts earn almost nothing")
add_bullet(doc, "Batch-produce: 1 long-form per week + 2–3 Shorts from each video")

add_heading(doc, "🏆 Top 5 Video Ideas", 2, "1565C0")
add_video_idea(doc, 1, "📊",
    '"Trump\'s Tariffs Will Cost You $700 This Year — Here\'s How to Get It Back"',
    "Ties a breaking news trigger to actionable personal finance. High search volume, emotional stakes, solves a real problem.",
    "Talking head + screen graphics, 12–15 min",
    "12–15 minutes",
    '"Everyone is talking about tariffs. Nobody is telling you what to actually DO with your money right now."'
)
add_video_idea(doc, 2, "💡",
    '"The Passive Income Lie Nobody\'s Talking About (And What Actually Works in 2026)"',
    "Taps massive passive income search volume while differentiating as the honest voice. Speaks to 'Richer By Design' ethos.",
    "Story-driven talking head, 15–18 min",
    "15–18 minutes",
    '"I made $4,700 in so-called passive income last year. After expenses? Negative $200. Here\'s what nobody tells you."'
)
add_video_idea(doc, 3, "⏰",
    '"The Fed Is Cutting Rates in March — Move Your Money NOW Before It\'s Too Late"',
    "Urgency + specific timing + FOMO. Real event happening NOW (March 2026). Searchable, timely, actionable.",
    "Voiceover with charts + talking head outro, 10–12 min",
    "10–12 minutes",
    '"If you have money in a regular savings account right now, you are actively losing purchasing power. Here\'s exactly what to do."'
)
add_video_idea(doc, 4, "🎨",
    '"Stop Saving, Start Designing: The Wealth Blueprint Nobody Teaches You"',
    "Positions 'Richer By Design' as a philosophy. Appeals to people tired of generic budgeting advice. Builds brand identity.",
    "Talking head, cinematic/aesthetic visuals, 15–20 min",
    "15–20 minutes",
    '"Every financial guru says: save more, spend less. Nobody ever taught you how to actually design a wealthy life."'
)
add_video_idea(doc, 5, "🤖",
    '"AI Is Changing How the Rich Invest — Here\'s How to Do It With $1,000"',
    "Combines two massive trends (AI + personal finance). Democratizes topic for average viewers. Strong affiliate potential.",
    "Screen share + talking head, 12–15 min",
    "12–15 minutes",
    '"The wealthiest investors in the world are quietly using AI to make decisions in seconds that used to take weeks."'
)

add_heading(doc, "📈 90-Day Growth Tips", 2, "1565C0")
add_bullet(doc, "Own a unique framework — create a proprietary system like 'The Design Wealth Method' to build a signature brand identity")
add_bullet(doc, "Post tariff/rate cut content ASAP — topical content within the news cycle captures trending search traffic")
add_bullet(doc, "Finance CPM is gold ($12–$22) — avoid mixing off-niche content to protect your ad revenue rate")
add_bullet(doc, "Build an email list from Day 1 — finance audiences have high willingness to subscribe to newsletters")
add_bullet(doc, "Collab with adjacent creators in productivity, minimalism, and entrepreneurship to expand reach")

add_divider(doc)
doc.add_page_break()

# ============================================================
# CHANNEL 3: ARIA'S GREAT ADVENTURES
# ============================================================
add_heading(doc, "🌍 Channel 3: Aria's Great Adventures", 1, "E65100")
add_heading(doc, "Niche: Kids & Family Adventure | Edutainment | Live-Action Exploration", 2, "F57C00")

add_heading(doc, "🔥 Top Trends Right Now (March 2026)", 2, "E65100")
trends3 = [
    ("'Edutainment' Adventure Content is King", "Content combining real outdoor exploration with learning (nature facts, science, geography) shows highest retention rates."),
    ("Comfort & Familiarity Beats Viral Chasing", "Kids rewatch content obsessively. Channels with recognizable characters, consistent tone, and predictable episode structure outperform trend-chasers."),
    ("Parent-Safe = Algorithm Safe", "80% of parents with kids under 11 watch YouTube with them. Channels signaling trustworthiness get more algorithmic support and family device placement."),
    ("Spanish + Multilingual Content Exploding", "5 of the top 30 fastest-growing kids channels in 2026 are Spanish-language. Adding subtitles/dubs drives major international growth."),
    ("STEM/Science Adventures Surging", "Coding basics, nature experiments, and interactive phonics content pulling millions of views in March 2026 specifically."),
    ("Urban/Nature Exploration Gap", "City wildlife, backyard science, local nature walks for kids — hyper-local discovery content has loyal audiences and very low competition."),
]
for name, desc in trends3:
    add_bullet(doc, f"{desc}", f"{name}: ")

add_heading(doc, "🎬 Best Format", 2, "E65100")
add_bullet(doc, "Live-action adventure vlog with Aria as the lead — creates personal connection and authenticity (NOT animation)")
add_bullet(doc, "Supplement with simple animated graphics/overlays for facts — the 'Bluey effect' (fun for kids AND parents)")
add_bullet(doc, "Post 3–5 Shorts/week (teaser moments from adventures) + 2–3 long-form vlogs per week")
add_bullet(doc, "Launch a recurring named series immediately: 'Aria's World Wonders,' 'Aria's Science Adventures'")
add_bullet(doc, "Target the 6–10 age group specifically — sweet spot for live-action adventure content")

add_heading(doc, "🏆 Top 5 Video Ideas", 2, "E65100")
add_video_idea(doc, 1, "🕵️",
    '"We Found a SECRET Underground World Kids Can Explore!! 😱 (Family Adventure Vlog)"',
    "'SECRET' + 'underground' triggers curiosity gap clicks. Family vlog tag pulls in parent viewers. High search volume for kids adventure destinations.",
    "10–15 min long-form + 60s Short teaser of the 'big reveal'",
    "10–15 minutes",
    '"We were just hiking a normal trail when Aria spotted something NOBODY expected to find..."'
)
add_video_idea(doc, 2, "🦎",
    '"Aria Discovers 5 WILD Animals in Our Backyard (You Won\'t Believe #3!) 🐍🦎"',
    "Number list titles get 35% higher CTR. 'You Won't Believe' triggers completion. Local/backyard angle is relatable for all families.",
    "8–12 min educational adventure + daily Short of each animal reveal",
    "8–12 minutes",
    '"Aria, what\'s that moving in the grass?! Wait... is that a...?!"'
)
add_video_idea(doc, 3, "🏆",
    '"EXTREME Kids Obstacle Course Challenge in the JUNGLE!! 🌴 (Last One Wins!)"',
    "'Challenge,' 'Extreme,' and 'Last One Wins' are proven click triggers. Competitive element keeps kids watching to see the outcome.",
    "12–18 min high energy vlog + Short of hardest obstacle",
    "12–18 minutes",
    '"Aria is taking on the world\'s toughest jungle obstacle course — and ONLY one kid can win!"'
)
add_video_idea(doc, 4, "🎢",
    '"We Spent 24 HOURS at the WORLD\'S BIGGEST Kids\' Theme Park! (Aria\'s Reaction 😭)"',
    "'24 Hours' format is evergreen and highly searchable. Emotional reaction shot in title drives clicks. Theme park content perennially top-performing.",
    "15–20 min full-day vlog + multiple Shorts (rides, food, surprises)",
    "15–20 minutes",
    '"We told Aria we were going grocery shopping... then THIS happened."'
)
add_video_idea(doc, 5, "🌋",
    '"Aria Teaches Kids About VOLCANOES… Then Makes One ERUPT! 🌋 (Science Adventure Ep. 1)"',
    "Episodic label ('Ep. 1') builds return viewers. 'Teaches + shows' hits the edutainment sweet spot parents want.",
    "Series anchor — 10 min + 30s Short of the eruption moment",
    "10–12 minutes",
    '"Today, Aria is your science teacher... and class is about to get EXPLOSIVE!"'
)

add_heading(doc, "📈 90-Day Growth Tips", 2, "E65100")
add_bullet(doc, "Brand Aria as THE Explorer — give her a signature item (explorer hat, backpack, catchphrase). Kids watch FOR Aria, not just the topic.")
add_bullet(doc, "COPPA compliance is non-negotiable — mark appropriate videos as 'Made for Kids' to avoid demonetization")
add_bullet(doc, "Thumbnail formula: close-up of Aria's face showing strong emotion + bold 3-word text + vivid background")
add_bullet(doc, "End every video with a question ('Where should Aria explore next?') to drive comment engagement")
add_bullet(doc, "The 'girl-led explorer' niche is UNTAPPED — there is NO dominant girl-led kids adventure channel at scale. This is Aria's edge.")
add_bullet(doc, "Cross-promote on Instagram Reels — kids content creators are seeing 40–60% additional reach at zero extra production cost")

add_divider(doc)
doc.add_page_break()

# ============================================================
# CHANNEL 4: DASHON McFARLANE (Personal Brand)
# ============================================================
add_heading(doc, "🌌 Channel 4: Dashon McFarlane", 1, "4A148C")
add_heading(doc, "Niche: Personal Branding, Entrepreneurship & Marketing", 2, "6A1B9A")

add_heading(doc, "🔥 Top Trends Right Now (March 2026)", 2, "4A148C")
trends4 = [
    ("YouTube Shorts Hit 90 Billion Views/Day", "Short-form is the #1 discovery engine for new channels in 2026. Non-negotiable for growth from zero."),
    ("Quality > Quantity Era", "Daily posting is dead. Fewer, better videos with strong hooks and SEO optimization consistently win."),
    ("Serialized Content Exploding", "57% of consumers want creator SERIES, not one-offs. Episodic content drives return viewers and subscriber loyalty."),
    ("Authentic > Polished", "Raw, talking-head, unscripted content outperforms high-gloss production. Real beats manufactured."),
    ("AI Content Backlash", "46% of audiences are uncomfortable with AI influencers. Demand for REAL human voices has never been higher."),
    ("YouTube = New Google for Gen Z", "50%+ of Gen Z searches YouTube instead of Google. SEO-optimized video titles are more important than ever."),
    ("The Contrarian Hook Dominates", "Videos titled 'Stop doing X' / 'The old way is DEAD' are spiking CTR across all creator niches."),
]
for name, desc in trends4:
    add_bullet(doc, f"{desc}", f"{name}: ")

add_heading(doc, "🎬 Best Format", 2, "4A148C")
add_bullet(doc, "Go Talking Head — personal branding is LITERALLY about you. Your face, voice, and energy ARE the brand.")
add_bullet(doc, "Use Shorts as trailers — post 30–60 sec clips from each long-form video 24hrs before the full video")
add_bullet(doc, "Launch a content SERIES ('Dashon's Brand Breakdown' or 'Marketing Mondays') for appointment viewing")
add_bullet(doc, "Production tip: Clean background + good lighting + iPhone + Rode mic = 80% of the quality you need")
add_bullet(doc, "Cadence: 1 long-form (12–20 min) per week + 3 Shorts per week")

add_heading(doc, "🏆 Top 5 Video Ideas", 2, "4A148C")
add_video_idea(doc, 1, "📈",
    '"I Built a Personal Brand From 0 to 10K Subscribers in 90 Days — Here\'s My Exact Strategy"',
    "Social proof + specific numbers + time frame = irresistible click. Search-optimized for 'how to grow YouTube channel' queries.",
    "Talking head + screen recording/B-roll, 14–18 min",
    "14–18 minutes",
    '"Most people spend 2 years trying to grow on YouTube. I figured out how to do it in 90 days — including what FAILED."'
)
add_video_idea(doc, 2, "🔥",
    '"The Personal Branding Strategy EVERYONE Is Teaching Is Completely Wrong (Do This Instead)"',
    "Contrarian hooks are the #1 performing title format in 2026. Triggers curiosity and challenges existing beliefs.",
    "Talking head, high energy, fast cuts, 10–14 min",
    "10–14 minutes",
    '"Every guru tells you: post every day, pick a niche, be consistent. But what if that\'s exactly why you\'re stuck at 200 subscribers?"'
)
add_video_idea(doc, 3, "🤖",
    '"These 7 AI Tools Are Making Personal Brand Creators Rich in 2026 (I Tested All of Them)"',
    "AI is the hottest search category on YouTube. Combining it with personal brand monetization = massive search volume.",
    "Screen recording tutorial + talking head intro/outro, 15–20 min",
    "15–20 minutes",
    '"The people winning online right now aren\'t working harder — they\'ve plugged into these 7 tools that automate 80% of their content workflow."'
)
add_video_idea(doc, 4, "🎯",
    '"How to Start a YouTube Channel in 2026 With NO Audience, NO Equipment, and NO Experience (Step-by-Step)"',
    "'How to start a YouTube channel' is a perennial top-10 search term. The 2026 angle + NO [obstacle] formula makes it hyper-clickable.",
    "Structured tutorial, talking head with clear sections, 18–25 min",
    "18–25 minutes",
    '"You don\'t need a ring light. You don\'t need a fancy camera. You don\'t need a big following. Here\'s the actual truth about starting from zero."'
)
add_video_idea(doc, 5, "💀",
    '"Why 99% of Personal Brand Channels FAIL in Their First Year (And the 1 Shift That Changes Everything)"',
    "Fear-based framing + promise of 'the secret' = top CTR formula. Aligns with self-improvement + entrepreneurship crossover audience.",
    "Storytelling talking head, vulnerability-forward, 12–16 min",
    "12–16 minutes",
    '"I studied 100 personal brand channels that died before hitting 1,000 subscribers. Every single one made the same mistake."'
)

add_heading(doc, "📈 90-Day Growth Tips", 2, "4A148C")
add_bullet(doc, "Niche down ruthlessly — 'Entrepreneurship' is too broad. Own 'personal brand YouTube strategy for first-gen entrepreneurs' or similar")
add_bullet(doc, "SEO first — treat your channel like a search engine. Every video must answer a real question people type into YouTube")
add_bullet(doc, "Comment on 10 videos/day in your niche — real engagement builds visibility and community")
add_bullet(doc, "Custom thumbnails on EVERY video — they drive 154% more clicks than auto-generated")
add_bullet(doc, "Repurpose every long-form into: 3 Shorts, 1 LinkedIn post, 1 Twitter/X thread")
add_bullet(doc, "YouTube now allows up to 5 channels to co-author a video — use this for collaborations")
add_bullet(doc, "Track Watch Time + Click-Through Rate (CTR), NOT subscriber count — these drive the algorithm")

add_divider(doc)
doc.add_page_break()

# ============================================================
# MASTER CONTENT CALENDAR
# ============================================================
add_heading(doc, "📅 90-Day Master Launch Calendar", 1, "1A1A2E")
add_body(doc, "Use this schedule across all 4 channels to build momentum simultaneously.")
add_body(doc, "")

add_heading(doc, "Month 1 — Foundation (Weeks 1–4)", 2, "1A1A2E")
add_bullet(doc, "Set up all 4 channels correctly: banner art, About section, channel trailer, featured video")
add_bullet(doc, "Post the #1 Video Idea for each channel in Week 1 (your 'anchor' video)")
add_bullet(doc, "Launch daily Shorts on ALL channels from Day 1 — even 30-second clips work")
add_bullet(doc, "NaturalHealingRevival: Start the 'Herb of the Week' short-form series")
add_bullet(doc, "Richer By Design: Drop the tariff video IMMEDIATELY while it's in the news cycle")
add_bullet(doc, "Aria's Great Adventures: Launch 'Aria's World Wonders — Episode 1' as the series anchor")
add_bullet(doc, "Dashon McFarlane: Post your origin story / personal brand manifesto video")

add_heading(doc, "Month 2 — Discovery (Weeks 5–8)", 2, "1A1A2E")
add_bullet(doc, "3 Shorts/week per channel + 1 long-form/week per channel = 12 Shorts + 4 long-forms total")
add_bullet(doc, "Optimize every title with YouTube search intent (use VidIQ or TubeBuddy)")
add_bullet(doc, "NaturalHealingRevival: Launch '30-Day Herbal Reset Challenge' series")
add_bullet(doc, "Richer By Design: Post the 'Wealth Design Blueprint' brand cornerstone video")
add_bullet(doc, "Aria's Great Adventures: Add Spanish subtitles to top 5 videos for international reach")
add_bullet(doc, "Dashon McFarlane: Launch 'Marketing Mondays' weekly series")
add_bullet(doc, "Start cross-promoting: mention other channels in videos where relevant")

add_heading(doc, "Month 3 — Authority (Weeks 9–12)", 2, "1A1A2E")
add_bullet(doc, "Increase Shorts to 5/week per channel using repurposed long-form clips")
add_bullet(doc, "Identify top 3 performing videos per channel and double down on similar topics")
add_bullet(doc, "Begin outreach for collaborations in adjacent niches")
add_bullet(doc, "NaturalHealingRevival: Add affiliate links for herbs/supplements (high-converting in health niche)")
add_bullet(doc, "Richer By Design: Launch email newsletter / lead magnet from YouTube CTA")
add_bullet(doc, "Aria's Great Adventures: Community post asking 'Where should Aria go next?' to drive votes/engagement")
add_bullet(doc, "Dashon McFarlane: Publish 'Results after 90 Days' transparency video (builds massive trust)")

add_divider(doc)
doc.add_page_break()

# ============================================================
# FINAL TIPS & UNIVERSAL RULES
# ============================================================
add_heading(doc, "⚡ Universal Rules for All 4 Channels", 1, "B71C1C")

rules = [
    ("Thumbnails Are Everything", "Your thumbnail is a billboard. Close-up face with strong emotion + bold 3-word text + vivid color = highest CTR. Test 3 per video in the first week."),
    ("Post Consistently, Not Constantly", "1 quality long-form + 3–5 Shorts per week beats daily mediocre uploads every time in 2026."),
    ("Shorts = Discovery, Long-Form = Loyalty", "Shorts get new eyes on your channel. Long-form converts those viewers into subscribers and earns real ad revenue."),
    ("SEO Your Titles Before Posting", "YouTube IS a search engine. Every title must answer a question people are actually searching. Use VidIQ or TubeBuddy free plans."),
    ("The First 30 Seconds Win or Lose Everything", "Hook viewers immediately. State the promise, the pain, or the payoff in the first 30 seconds — or they're gone."),
    ("Engagement Signals Drive the Algorithm", "Comments > Likes > Views for the algorithm. End every video with a specific question that demands a response."),
    ("Be Patient — Channels Usually Explode Around Video 20–30", "Most channels hit their first viral video between video 20 and 40. Don't quit before the breakthrough."),
]

for name, desc in rules:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r1 = p.add_run(f"✅ {name}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

add_divider(doc)

# Footer note
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.paragraph_format.space_before = Pt(20)
rf = footer_para.add_run("YouTube Growth Master Plan — Generated by MaxClaw AI | March 2026\nResearch based on live YouTube, SEO, and industry data from March 2026.")
rf.font.size = Pt(9)
rf.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# Save
doc.save("/workspace/YouTube_Growth_Master_Plan_March2026.docx")
print("SUCCESS: Document saved to /workspace/YouTube_Growth_Master_Plan_March2026.docx")
